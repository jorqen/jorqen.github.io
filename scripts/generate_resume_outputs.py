#!/usr/bin/env python3
"""Generate website data and downloadable resume files from one YAML source."""

from __future__ import annotations

import argparse
from calendar import monthrange
from datetime import date
from functools import lru_cache
import hashlib
import html
import json
import shutil
import subprocess
import sys
import textwrap
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit

try:
    import yaml
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing YAML dependency. Install PyYAML or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing DOCX dependencies. Install python-docx or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing PDF dependencies. Install reportlab or run through scripts/build_resume_formats.sh."
    ) from exc

ROOT = Path(__file__).resolve().parents[1]

MONTH_NAMES = {
    "en": ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"],
    "ru": ["Янв", "Фев", "Мар", "Апр", "Май", "Июн", "Июл", "Авг", "Сен", "Окт", "Ноя", "Дек"],
}

DOWNLOAD_STYLES = {
    "colors": {
        "accent": "1D5FD1",
        "accentMuted": "EAF2FF",
        "surface": "F6F9FD",
        "surfaceAlt": "EEF4FB",
        "text": "0F1F36",
        "muted": "526581",
        "line": "DBE4F2",
        "white": "FFFFFF",
    },
    "typography": {
        "kicker": {"fontSize": 8.5, "lineHeight": 10},
        "title": {"fontSize": 24, "lineHeight": 28},
        "headline": {"fontSize": 10.8, "lineHeight": 13},
        "contact": {"fontSize": 8.2, "lineHeight": 10.2},
        "summary": {"fontSize": 9.2, "lineHeight": 12.1},
        "factLabel": {"fontSize": 7.2, "lineHeight": 8.6},
        "factValue": {"fontSize": 8.6, "lineHeight": 10.8},
        "section": {"fontSize": 11.3, "lineHeight": 13.3},
        "sectionBody": {"fontSize": 8.5, "lineHeight": 10.8},
        "company": {"fontSize": 10.4, "lineHeight": 12.3},
        "meta": {"fontSize": 8.7, "lineHeight": 10.8},
        "body": {"fontSize": 9, "lineHeight": 11.8},
        "bullet": {"fontSize": 8.8, "lineHeight": 11},
        "cardTitle": {"fontSize": 8.8, "lineHeight": 10.8},
        "cardBody": {"fontSize": 8.4, "lineHeight": 10.6},
    },
}

DOWNLOAD_DOCX_PAGE_MARGINS = {
    "horizontal": 0.50,
    "vertical": 0.45,
}

DOWNLOAD_PDF_PAGE_MARGINS = {
    "horizontal": 0.42,
    "vertical": 0.37,
}

PDF_TYPOGRAPHY_SCALE = 0.9

SITE_UI = {
    "metaDescription": {
        "en": "Matvey Sizov - Backend Developer / Software Engineer. Go, distributed systems, low-latency backend development, and production reliability in product and platform teams.",
        "ru": "Матвей Сизов - Backend Developer / Software Engineer. Go, распределенные системы, backend с низкой задержкой и надежность production в продуктовых и платформенных командах.",
    },
    "navResume": {
        "en": "Resume",
        "ru": "Резюме",
    },
    "langSwitcherLabel": {
        "en": "Language switcher",
        "ru": "Переключение языка",
    },
    "theme": {
        "toDark": {
            "en": "Dark theme",
            "ru": "Тёмная тема",
        },
        "toLight": {
            "en": "Light theme",
            "ru": "Светлая тема",
        },
        "switcherLabel": {
            "en": "Theme switcher",
            "ru": "Переключение темы",
        },
    },
    "resumeDownloads": {
        "title": {
            "en": "Download resume",
            "ru": "Скачать резюме",
        },
        "labels": {
            "pdf": {
                "en": "For sharing and printing",
                "ru": "Для отправки и печати",
            },
            "docx": {
                "en": "Editable source",
                "ru": "Редактируемый источник",
            },
            "txt": {
                "en": "Text version",
                "ru": "Текстовая версия",
            },
        },
    },
    "lightbox": {
        "openPhoto": {
            "en": "Open photo",
            "ru": "Открыть фото",
        },
        "close": {
            "en": "Close photo viewer",
            "ru": "Закрыть просмотр фото",
        },
        "previous": {
            "en": "Previous photo",
            "ru": "Предыдущее фото",
        },
        "next": {
            "en": "Next photo",
            "ru": "Следующее фото",
        },
    },
    "footer": {
        "en": "© {year} {name}. Personal website for recruiters and hiring managers.",
        "ru": "© {year} {name}. Сайт-визитка для рекрутеров и менеджеров по найму.",
    },
}

DOWNLOAD_SECTION_TITLES = {
    "en": {
        "profile": "ABOUT ME",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "skills": "PROGRAMMING SKILLS",
    },
    "ru": {
        "profile": "ОБО МНЕ",
        "experience": "ПРОФЕССИОНАЛЬНЫЙ ОПЫТ",
        "education": "ОБРАЗОВАНИЕ",
        "skills": "ТЕХНИЧЕСКИЕ НАВЫКИ",
    },
}

HERO_CONTACT_KEYS = ("linkedin", "github", "telegram")


def localized_tree(value: Any, lang: str) -> Any:
    if isinstance(value, dict) and set(value.keys()).issubset({"en", "ru"}):
        return value[lang]
    if isinstance(value, dict):
        return {key: localized_tree(item, lang) for key, item in value.items()}
    if isinstance(value, list):
        return [localized_tree(item, lang) for item in value]
    return value


def load_yaml_mapping(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def hex_color(value: str) -> str:
    return str(value).strip().removeprefix("#").upper()


def rgb_color(hex_value: str) -> RGBColor:
    color = hex_color(hex_value)
    return RGBColor(
        int(color[0:2], 16),
        int(color[2:4], 16),
        int(color[4:6], 16),
    )


def download_styles() -> dict[str, Any]:
    return {
        "colors": {
            key: hex_color(value)
            for key, value in DOWNLOAD_STYLES["colors"].items()
        },
        "typography": {
            key: value.copy()
            for key, value in DOWNLOAD_STYLES["typography"].items()
        },
    }


def parse_resume_date(value: str, *, upper_bound: bool) -> date:
    parts = value.split("-")
    year = int(parts[0])
    if len(parts) == 1:
        month = 12 if upper_bound else 1
        day = 31 if upper_bound else 1
        return date(year, month, day)
    month = int(parts[1])
    day = monthrange(year, month)[1] if upper_bound else 1
    return date(year, month, day)


def months_between(start_date: date, end_date: date) -> int:
    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
    if end_date.day < start_date.day:
        months -= 1
    return max(months, 0)


def format_duration_label(total_months: int, lang: str) -> str:
    years, months = divmod(total_months, 12)
    parts: list[str] = []
    if years:
        if lang == "ru":
            parts.append(f"{years} г.")
        else:
            parts.append(f"{years} {'yr' if years == 1 else 'yrs'}")
    if months or not parts:
        if lang == "ru":
            parts.append(f"{months} мес.")
        else:
            parts.append(f"{months} {'mo' if months == 1 else 'mos'}")
    return " ".join(parts)


def experience_duration(item: dict[str, Any], lang: str) -> str:
    end_date = parse_resume_date(item["endDate"], upper_bound=True) if item.get("endDate") else date.today()
    start_date = parse_resume_date(item["startDate"], upper_bound=False)
    return format_duration_label(months_between(start_date, end_date), lang)


def contact_value(contact: dict[str, str]) -> str:
    return str(contact["value"]).strip()


def is_web_url(value: str) -> bool:
    parsed = urlsplit(value)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def contact_href(contact: dict[str, str]) -> str | None:
    value = contact_value(contact)
    return value if is_web_url(value) else None


def contact_link_text(contact: dict[str, str], *, shorten_web_urls: bool = True) -> str:
    value = contact_value(contact)
    if shorten_web_urls and is_web_url(value):
        return value.removeprefix("https://").removeprefix("http://")
    return value


def convert_svg_to_png(source: Path, build_dir: Path, size: int = 48) -> Path | None:
    rsvg = shutil.which("rsvg-convert")
    if not rsvg:
        return None
    build_dir.mkdir(parents=True, exist_ok=True)
    relative_source = source.relative_to(ROOT)
    cache_key = hashlib.sha1(f"{relative_source}:preserve-aspect-v2".encode("utf-8")).hexdigest()[:10]
    output = build_dir / f"{source.stem}-{cache_key}-{size}.png"
    if output.exists() and output.stat().st_mtime >= source.stat().st_mtime:
        return output
    subprocess.run(
        [rsvg, "-w", str(size), "-o", str(output), str(source)],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return output


def resolve_image(path_value: str | None, build_dir: Path, size: int = 48) -> Path | None:
    if not path_value:
        return None
    source = ROOT / path_value
    if not source.exists():
        return None
    if source.suffix.lower() == ".svg":
        return convert_svg_to_png(source, build_dir, size)
    if source.suffix.lower() in {".png", ".jpg", ".jpeg"}:
        return source
    return None


def contact_icon_path(contact_key: str, contact: dict[str, str], build_dir: Path) -> Path | None:
    icon_path = contact.get("icon") or contact.get("iconDark")
    if not icon_path and contact_key == "website":
        icon_path = "assets/icons/website.svg"
    return resolve_image(icon_path, build_dir, 32)


def company_icon_path(item: dict[str, Any], build_dir: Path) -> Path | None:
    return resolve_image(item.get("companyIcon") or item.get("companyIconDark"), build_dir, 40)


def institution_icon_path(item: dict[str, Any], build_dir: Path) -> Path | None:
    return resolve_image(item.get("institutionIcon") or item.get("institutionIconDark"), build_dir, 40)


def download_section_title(lang: str, key: str) -> str:
    return DOWNLOAD_SECTION_TITLES[lang][key]


def experience_header_text(item: dict[str, Any]) -> str:
    return f"{item['role']} ({item['location']}, {item['period']}, {item['duration']})"


def education_header_text(item: dict[str, Any]) -> str:
    return f"{item['degree']} ({item['period']})"


def skill_group_segments(groups: list[dict[str, Any]]) -> list[tuple[str, str]]:
    return [(group["title"], ", ".join(group["items"])) for group in groups]


def format_resume_date(value: str, month_names: list[str]) -> str:
    parts = value.split("-")
    if len(parts) == 1:
        return parts[0]
    year, month = parts
    return f"{month_names[int(month) - 1]} {year}"


def format_period(item: dict[str, Any], labels: dict[str, str], month_names: list[str]) -> str:
    start = format_resume_date(item["startDate"], month_names)
    end_date = item.get("endDate")
    end = format_resume_date(end_date, month_names) if end_date else labels["present"]
    return f"{start} - {end}"


def resume_date_upper_bound(value: str | None) -> date | None:
    if value is None:
        return None
    return parse_resume_date(value, upper_bound=True)


def is_expected_education(item: dict[str, Any]) -> bool:
    return resume_date_upper_bound(item["endDate"]) > date.today()


def format_education_period(item: dict[str, Any], labels: dict[str, str], month_names: list[str]) -> str:
    start = format_resume_date(item["startDate"], month_names)
    end_date = format_resume_date(item["endDate"], month_names)
    end = f"{labels['expectedGraduation']}: {end_date}" if is_expected_education(item) else end_date
    return f"{start} - {end}"


def download_file_names(file_base_name: str) -> dict[str, str]:
    return {
        "pdf": f"{file_base_name}.pdf",
        "docx": f"{file_base_name}.docx",
        "txt": f"{file_base_name}.txt",
    }


def resume_file_base_name(source: dict[str, Any]) -> str:
    english_name = localized_tree(source["person"]["name"], "en")
    return "_".join(english_name.split())


def site_meta(person: dict[str, Any], lang: str) -> dict[str, str]:
    return {
        "title": f"{person['name']} | {person['headline']}",
        "description": SITE_UI["metaDescription"][lang],
    }


def site_nav(
    labels: dict[str, str],
    experience: dict[str, Any],
    education: dict[str, Any],
    strengths: dict[str, Any],
    lang: str,
) -> dict[str, str]:
    return {
        "resume": SITE_UI["navResume"][lang],
        "experience": experience["title"],
        "education": education["title"],
        "strengths": strengths["title"],
        "skills": labels["stack"],
    }


def site_footer(person: dict[str, Any], lang: str) -> str:
    return SITE_UI["footer"][lang].format(year="{year}", name=person["name"])


def language_content(source: dict[str, Any], lang: str) -> dict[str, Any]:
    person = localized_tree(source["person"], lang)
    labels = localized_tree(source["resumeLabels"], lang)
    month_names = MONTH_NAMES[lang]

    experience = localized_tree(source["experience"], lang)
    experience_items = []
    optional_experience_fields = (
        ("companyIcon", "icon"),
        ("companyIconDark", "iconDark"),
        ("companyUrl", "url"),
    )
    for item in sorted(
        experience["items"],
        key=lambda experience_item: resume_date_upper_bound(experience_item["startDate"]),
        reverse=True,
    ):
        output_item = {
            "company": item["company"],
            "role": item["role"],
            "period": format_period(item, labels, month_names),
            "duration": experience_duration(item, lang),
            "location": item["location"],
            "intro": item["summary"],
            "bullets": item["highlights"],
            "stack": item["stack"],
        }
        output_item.update({
            output_key: item[source_key]
            for output_key, source_key in optional_experience_fields
            if source_key in item
        })
        experience_items.append(output_item)
    experience["items"] = experience_items

    education = localized_tree(source["education"], lang)
    education_items = []
    optional_education_fields = (
        ("institutionIcon", "icon"),
        ("institutionIconDark", "iconDark"),
        ("institutionUrl", "url"),
    )
    for item in education["items"]:
        output_item = {
            "institution": item["institution"],
            "degree": item["degree"],
            "period": format_education_period(item, labels, month_names),
        }
        output_item.update({
            output_key: item[source_key]
            for output_key, source_key in optional_education_fields
            if source_key in item
        })
        education_items.append(output_item)
    education["items"] = education_items
    strengths = localized_tree(source["strengths"], lang)

    resume = {
        "fileBaseName": resume_file_base_name(source),
        "title": SITE_UI["resumeDownloads"]["title"][lang],
        "labels": localized_tree(SITE_UI["resumeDownloads"]["labels"], lang),
    }
    file_names = download_file_names(resume["fileBaseName"])
    resume["files"] = {
        extension: f"resume/{lang}/{name}"
        for extension, name in file_names.items()
    }
    resume["downloadNames"] = file_names

    return {
        "meta": site_meta(person, lang),
        "brand": person["name"],
        "nav": site_nav(labels, experience, education, strengths, lang),
        "langSwitcherLabel": SITE_UI["langSwitcherLabel"][lang],
        "theme": localized_tree(SITE_UI["theme"], lang),
        "hero": {
            "kicker": person["headline"],
            "name": person["name"],
            "role": person["role"],
            "summary": person["summary"],
            "photo": person["photo"],
            "facts": person["facts"],
        },
        "resume": resume,
        "experience": experience,
        "education": education,
        "strengths": strengths,
        "skills": localized_tree(source["skills"], lang),
        "preferences": localized_tree(source["preferences"], lang),
        "gallery": localized_tree(source["gallery"], lang),
        "lightbox": localized_tree(SITE_UI["lightbox"], lang),
        "footer": site_footer(person, lang),
    }


def build_resume_data(source: dict[str, Any]) -> dict[str, Any]:
    languages = source["languages"]
    contact_keys = list(source["contacts"].keys())
    content = {lang: language_content(source, lang) for lang in languages}
    contacts = {
        key: {lang: localized_tree(contact, lang) for lang in languages}
        for key, contact in source["contacts"].items()
    }
    labels = {lang: localized_tree(source["resumeLabels"], lang) for lang in languages}
    return {
        "meta": {
            "schema": source["schema"],
            "defaultLanguage": source["defaultLanguage"],
            "languages": languages,
        },
        "contactKeys": contact_keys,
        "downloadStyles": download_styles(),
        "contacts": contacts,
        "labels": labels,
        "content": content,
    }


def load_data(path: Path) -> dict[str, Any]:
    source = load_yaml_mapping(path)
    return build_resume_data(source)


def contact_parts(data: dict[str, Any], lang: str) -> list[str]:
    parts = []
    for key in data["contactKeys"]:
        contact = data["contacts"][key][lang]
        parts.append(contact_link_text(contact, shorten_web_urls=False))
    return parts


def wrap_line(text: str, width: int = 100) -> list[str]:
    if not text:
        return []
    return textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=False) or [text]


def add_wrapped(lines: list[str], text: str = "", width: int = 100) -> None:
    if text:
        lines.extend(wrap_line(text, width))
    else:
        lines.append("")


def generate_txt(data: dict[str, Any], lang: str, output_path: Path) -> None:
    content = data["content"][lang]
    lines: list[str] = []

    lines.append(content["hero"]["name"])
    lines.append("=" * len(content["hero"]["name"]))
    lines.append(" | ".join(contact_parts(data, lang)))
    add_wrapped(lines)
    add_wrapped(lines, content["hero"]["role"])

    add_wrapped(lines)
    lines.append(download_section_title(lang, "profile"))
    lines.append("-" * len(download_section_title(lang, "profile")))
    add_wrapped(lines, content["hero"]["summary"], 100)

    add_wrapped(lines)
    lines.append(download_section_title(lang, "experience"))
    lines.append("-" * len(download_section_title(lang, "experience")))
    for item in content["experience"]["items"]:
        add_wrapped(lines)
        lines.append(f"{item['company']} | {experience_header_text(item)}")
        if item.get("companyUrl"):
            lines.append(f"{content['experience']['companySiteLabel']}: {item['companyUrl']}")
        add_wrapped(lines, item["intro"])
        for bullet in item["bullets"]:
            for index, wrapped in enumerate(wrap_line(bullet, 96)):
                prefix = "- " if index == 0 else "  "
                lines.append(f"{prefix}{wrapped}")
        lines.append(f"{data['labels'][lang]['stack']}: {', '.join(item['stack'])}")

    add_wrapped(lines)
    lines.append(download_section_title(lang, "education"))
    lines.append("-" * len(download_section_title(lang, "education")))
    for item in content["education"]["items"]:
        add_wrapped(lines)
        lines.append(f"{item['institution']} | {education_header_text(item)}")
        education_site_label = content["education"].get("institutionSiteLabel")
        if education_site_label and item.get("institutionUrl"):
            lines.append(f"{education_site_label}: {item['institutionUrl']}")

    add_wrapped(lines)
    lines.append(download_section_title(lang, "skills"))
    lines.append("-" * len(download_section_title(lang, "skills")))
    for title, values in skill_group_segments(content["skills"]["groups"]):
        lines.append(f"{title}: {values}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def add_hyperlink(paragraph: Any, text: str, url: str, color_hex: str, bold: bool = False) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(fonts)
    properties.append(color)
    properties.append(underline)
    if bold:
        properties.append(OxmlElement("w:b"))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # noqa: SLF001 - python-docx exposes no public hyperlink writer.


def docx_paragraph_style(styles: Any, name: str) -> Any:
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_docx_style_font_family(style: Any, font_name: str) -> None:
    properties = style._element.get_or_add_rPr()  # noqa: SLF001
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.append(fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def set_docx_style_widow_control(style: Any, enabled: bool = False) -> None:
    paragraph_properties = style._element.get_or_add_pPr()  # noqa: SLF001
    widow_control = paragraph_properties.find(qn("w:widowControl"))
    if widow_control is None:
        widow_control = OxmlElement("w:widowControl")
        paragraph_properties.append(widow_control)
    widow_control.set(qn("w:val"), "1" if enabled else "0")


def configure_docx_paragraph_style(
    styles: Any,
    name: str,
    font_name: str,
    font_size: float,
    color: RGBColor,
    *,
    bold: bool = False,
    alignment: Any = None,
    keep_with_next: bool = False,
    line_spacing: float = 12,
    space_before: float = 0,
    space_after: float = 0,
    left_indent: float | None = None,
    first_line_indent: float | None = None,
) -> None:
    style = docx_paragraph_style(styles, name)
    style.font.name = font_name
    set_docx_style_font_family(style, font_name)
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.alignment = alignment
    style.paragraph_format.keep_with_next = keep_with_next
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.line_spacing = Pt(line_spacing)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    set_docx_style_widow_control(style, enabled=False)
    if left_indent is not None:
        style.paragraph_format.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        style.paragraph_format.first_line_indent = Inches(first_line_indent)


def configure_docx(document: Document, download_style: dict[str, Any]) -> None:
    style_colors = download_style["colors"]
    typography = download_style["typography"]
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(DOWNLOAD_DOCX_PAGE_MARGINS["vertical"])
    section.bottom_margin = Inches(DOWNLOAD_DOCX_PAGE_MARGINS["vertical"])
    section.left_margin = Inches(DOWNLOAD_DOCX_PAGE_MARGINS["horizontal"])
    section.right_margin = Inches(DOWNLOAD_DOCX_PAGE_MARGINS["horizontal"])

    font_name = "Arial"
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    set_docx_style_font_family(normal, font_name)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = rgb_color(style_colors["text"])
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(11.8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    set_docx_style_widow_control(normal, enabled=False)

    configure_docx_paragraph_style(
        styles,
        "ResumeKicker",
        font_name,
        typography["kicker"]["fontSize"],
        rgb_color(style_colors["accent"]),
        bold=True,
        line_spacing=typography["kicker"]["lineHeight"],
        space_after=2,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeTitle",
        font_name,
        typography["title"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["title"]["lineHeight"],
        space_after=2,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeHeadline",
        font_name,
        typography["headline"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["headline"]["lineHeight"],
        space_after=3,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeContact",
        font_name,
        typography["contact"]["fontSize"],
        rgb_color(style_colors["muted"]),
        line_spacing=typography["contact"]["lineHeight"],
        space_after=3,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeSummary",
        font_name,
        typography["summary"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["summary"]["lineHeight"],
        space_after=0,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeFactLabel",
        font_name,
        typography["factLabel"]["fontSize"],
        rgb_color(style_colors["muted"]),
        bold=True,
        line_spacing=typography["factLabel"]["lineHeight"],
        space_after=1,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeFactValue",
        font_name,
        typography["factValue"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["factValue"]["lineHeight"],
        space_after=0,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeSection",
        font_name,
        typography["section"]["fontSize"],
        rgb_color(style_colors["accent"]),
        bold=True,
        line_spacing=typography["section"]["lineHeight"],
        space_before=8,
        space_after=2,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeSectionBody",
        font_name,
        typography["sectionBody"]["fontSize"],
        rgb_color(style_colors["muted"]),
        line_spacing=typography["sectionBody"]["lineHeight"],
        space_after=4,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeCompany",
        font_name,
        typography["company"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["company"]["lineHeight"],
        space_after=2,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeMeta",
        font_name,
        typography["meta"]["fontSize"],
        rgb_color(style_colors["muted"]),
        line_spacing=typography["meta"]["lineHeight"],
        space_after=1,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeBody",
        font_name,
        typography["body"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["body"]["lineHeight"],
        space_after=2,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeBullet",
        font_name,
        typography["bullet"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["bullet"]["lineHeight"],
        space_after=1,
        left_indent=0.17,
        first_line_indent=-0.11,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeCardTitle",
        font_name,
        typography["cardTitle"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["cardTitle"]["lineHeight"],
        space_after=2,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeCardBody",
        font_name,
        typography["cardBody"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["cardBody"]["lineHeight"],
        space_after=0,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeGap",
        font_name,
        1,
        rgb_color(style_colors["white"]),
        line_spacing=2,
        space_after=0,
    )


def docx_contact_line(
    paragraph: Any,
    data: dict[str, Any],
    lang: str,
    accent_color: str,
    build_dir: Path,
) -> None:
    for index, key in enumerate(data["contactKeys"]):
        contact = data["contacts"][key][lang]
        if index:
            paragraph.add_run("  |  ")
        icon_path = contact_icon_path(key, contact, build_dir / "contact-icons")
        if icon_path:
            paragraph.add_run().add_picture(str(icon_path), width=Inches(0.11))
            paragraph.add_run(" ")
        href = contact_href(contact)
        if href:
            add_hyperlink(paragraph, contact_link_text(contact), href, accent_color)
        else:
            paragraph.add_run(contact_link_text(contact))


def generate_docx(data: dict[str, Any], lang: str, output_path: Path, build_dir: Path) -> None:
    content = data["content"][lang]
    labels = data["labels"][lang]
    download_style = data["downloadStyles"]
    style_colors = download_style["colors"]
    document = Document()
    configure_docx(document, download_style)

    document.add_paragraph(content["hero"]["name"], style="ResumeTitle")
    contact_paragraph = document.add_paragraph(style="ResumeContact")
    docx_contact_line(contact_paragraph, data, lang, style_colors["accent"], build_dir)
    document.add_paragraph(content["hero"]["role"], style="ResumeHeadline")

    document.add_paragraph(download_section_title(lang, "profile"), style="ResumeSection")
    document.add_paragraph(content["hero"]["summary"], style="ResumeBody")

    document.add_paragraph(download_section_title(lang, "experience"), style="ResumeSection")
    for item in content["experience"]["items"]:
        header = document.add_paragraph(style="ResumeCompany")
        icon_path = company_icon_path(item, build_dir / "company-icons")
        if icon_path:
            header.add_run().add_picture(str(icon_path), width=Inches(0.12))
            header.add_run(" ")
        if item.get("companyUrl"):
            add_hyperlink(header, item["company"], item["companyUrl"], style_colors["accent"], bold=True)
        else:
            company_run = header.add_run(item["company"])
            company_run.bold = True
        header.add_run(f" | {experience_header_text(item)}").bold = True

        document.add_paragraph(item["intro"], style="ResumeBody")
        for bullet in item["bullets"]:
            document.add_paragraph(f"• {bullet}", style="ResumeBullet")
        stack = document.add_paragraph(style="ResumeMeta")
        stack.add_run(f"{labels['stack']}: ").bold = True
        stack.add_run(", ".join(item["stack"]))

    document.add_paragraph(download_section_title(lang, "education"), style="ResumeSection")
    for item in content["education"]["items"]:
        paragraph = document.add_paragraph(style="ResumeBody")
        icon_path = institution_icon_path(item, build_dir / "institution-icons")
        if icon_path:
            paragraph.add_run().add_picture(str(icon_path), width=Inches(0.12))
            paragraph.add_run(" ")
        if item.get("institutionUrl"):
            add_hyperlink(paragraph, item["institution"], item["institutionUrl"], style_colors["accent"], bold=True)
        else:
            paragraph.add_run(item["institution"]).bold = True
        paragraph.add_run(f" | {education_header_text(item)}")

    document.add_paragraph(download_section_title(lang, "skills"), style="ResumeSection")
    skills_paragraph = document.add_paragraph(style="ResumeBody")
    for index, (title, values) in enumerate(skill_group_segments(content["skills"]["groups"])):
        if index:
            skills_paragraph.add_run().add_break()
        skills_paragraph.add_run(f"{title}: ").bold = True
        skills_paragraph.add_run(values)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def find_font(candidates: list[str]) -> Path | None:
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return path
    fc_match = shutil.which("fc-match")
    if fc_match:
        for family in ("DejaVu Sans", "Arial Unicode MS", "Arial"):
            result = subprocess.run(
                [fc_match, "-f", "%{file}", family],
                check=False,
                text=True,
                capture_output=True,
            )
            if result.returncode == 0 and result.stdout.strip():
                path = Path(result.stdout.strip())
                if path.exists():
                    return path
    return None


def register_pdf_fonts(lang: str) -> tuple[str, str]:
    regular = find_font([
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ])
    bold = find_font([
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]) or regular

    if not regular:
        raise RuntimeError("Could not find a Unicode TrueType font for PDF generation.")

    if "ResumeRegular" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("ResumeRegular", str(regular)))
    if "ResumeBold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("ResumeBold", str(bold)))
    return "ResumeRegular", "ResumeBold"


def pdf_link(text: str, url: str, color_hex: str) -> str:
    return f'<a href="{html.escape(url, quote=True)}"><font color="#{color_hex}">{html.escape(text)}</font></a>'


def pdf_text(text: str) -> str:
    return html.escape(text).replace("\n", "<br/>")


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(pdf_text(text), style)


def pdf_markup(markup: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(markup, style)


@lru_cache(maxsize=None)
def image_dimensions(path: str) -> tuple[float, float] | None:
    try:
        width, height = ImageReader(path).getSize()
    except OSError:
        return None
    if width <= 0 or height <= 0:
        return None
    return float(width), float(height)


def pdf_icon_markup(path: Path | None, size: int = 9) -> str:
    if not path:
        return ""
    dimensions = image_dimensions(str(path))
    height = size
    if dimensions:
        intrinsic_width, intrinsic_height = dimensions
        height = size * intrinsic_height / intrinsic_width
    return (
        f'<img src="{html.escape(str(path), quote=True)}" '
        f'width="{size:.2f}" height="{height:.2f}" valign="middle"/>'
    )


def build_pdf_styles(lang: str, download_style: dict[str, Any]) -> dict[str, ParagraphStyle]:
    style_colors = download_style["colors"]
    typography = download_style["typography"]
    regular, bold = register_pdf_fonts(lang)
    base = getSampleStyleSheet()

    def font_size(style_key: str) -> float:
        return typography[style_key]["fontSize"] * PDF_TYPOGRAPHY_SCALE

    def line_height(style_key: str) -> float:
        return typography[style_key]["lineHeight"] * PDF_TYPOGRAPHY_SCALE

    return {
        "kicker": ParagraphStyle(
            "ResumeKicker",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=font_size("kicker"),
            leading=line_height("kicker"),
            textColor=colors.HexColor(f"#{style_colors['accent']}"),
            spaceAfter=2,
        ),
        "title": ParagraphStyle(
            "ResumeTitle",
            parent=base["Heading1"],
            fontName=bold,
            fontSize=font_size("title"),
            leading=line_height("title"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=2,
        ),
        "headline": ParagraphStyle(
            "ResumeHeadline",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=font_size("headline"),
            leading=line_height("headline"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=3,
        ),
        "contact": ParagraphStyle(
            "ResumeContact",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("contact"),
            leading=line_height("contact"),
            textColor=colors.HexColor(f"#{style_colors['muted']}"),
            spaceAfter=3,
        ),
        "summary": ParagraphStyle(
            "ResumeSummary",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("summary"),
            leading=line_height("summary"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=0,
        ),
        "factLabel": ParagraphStyle(
            "ResumeFactLabel",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=font_size("factLabel"),
            leading=line_height("factLabel"),
            textColor=colors.HexColor(f"#{style_colors['muted']}"),
            spaceAfter=1,
        ),
        "factValue": ParagraphStyle(
            "ResumeFactValue",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=font_size("factValue"),
            leading=line_height("factValue"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=0,
        ),
        "section": ParagraphStyle(
            "ResumeSection",
            parent=base["Heading2"],
            fontName=bold,
            fontSize=font_size("section"),
            leading=line_height("section"),
            textColor=colors.HexColor(f"#{style_colors['accent']}"),
            spaceBefore=8,
            spaceAfter=2,
        ),
        "sectionBody": ParagraphStyle(
            "ResumeSectionBody",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("sectionBody"),
            leading=line_height("sectionBody"),
            textColor=colors.HexColor(f"#{style_colors['muted']}"),
            spaceAfter=4,
        ),
        "company": ParagraphStyle(
            "ResumeCompany",
            parent=base["Heading3"],
            fontName=bold,
            fontSize=font_size("company"),
            leading=line_height("company"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=2,
        ),
        "meta": ParagraphStyle(
            "ResumeMeta",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("meta"),
            leading=line_height("meta"),
            textColor=colors.HexColor(f"#{style_colors['muted']}"),
            spaceAfter=1,
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("body"),
            leading=line_height("body"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("bullet"),
            leading=line_height("bullet"),
            leftIndent=12,
            firstLineIndent=-8,
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=1,
        ),
        "cardTitle": ParagraphStyle(
            "ResumeCardTitle",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=font_size("cardTitle"),
            leading=line_height("cardTitle"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=2,
        ),
        "cardBody": ParagraphStyle(
            "ResumeCardBody",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=font_size("cardBody"),
            leading=line_height("cardBody"),
            textColor=colors.HexColor(f"#{style_colors['text']}"),
            spaceAfter=0,
        ),
    }


def pdf_section(story: list[Any], title: str, styles: dict[str, ParagraphStyle], subtitle: str | None = None) -> None:
    story.append(Paragraph(html.escape(title.upper()), styles["section"]))
    if subtitle:
        story.append(pdf_paragraph(subtitle, styles["sectionBody"]))


def pdf_contact_markup(data: dict[str, Any], lang: str, accent_color: str, build_dir: Path) -> str:
    items = []
    for key in data["contactKeys"]:
        contact = data["contacts"][key][lang]
        contact_text = contact_link_text(contact)
        href = contact_href(contact)
        contact_markup = pdf_link(contact_text, href, accent_color) if href else pdf_text(contact_text)
        items.append(
            (
                f'{pdf_icon_markup(contact_icon_path(key, contact, build_dir / "contact-icons"))} '
                f"{contact_markup}"
            ).strip()
        )
    return " | ".join(items)


def generate_pdf(data: dict[str, Any], lang: str, output_path: Path, build_dir: Path) -> None:
    content = data["content"][lang]
    labels = data["labels"][lang]
    download_style = data["downloadStyles"]
    style_colors = download_style["colors"]
    styles = build_pdf_styles(lang, download_style)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=DOWNLOAD_PDF_PAGE_MARGINS["horizontal"] * inch,
        leftMargin=DOWNLOAD_PDF_PAGE_MARGINS["horizontal"] * inch,
        topMargin=DOWNLOAD_PDF_PAGE_MARGINS["vertical"] * inch,
        bottomMargin=DOWNLOAD_PDF_PAGE_MARGINS["vertical"] * inch,
        title=content["meta"]["title"],
        author=content["hero"]["name"],
    )
    story: list[Any] = []
    story.append(pdf_paragraph(content["hero"]["name"], styles["title"]))
    story.append(pdf_markup(pdf_contact_markup(data, lang, style_colors["accent"], build_dir), styles["contact"]))
    story.append(pdf_paragraph(content["hero"]["role"], styles["headline"]))

    pdf_section(story, download_section_title(lang, "profile"), styles)
    story.append(pdf_paragraph(content["hero"]["summary"], styles["body"]))

    pdf_section(story, download_section_title(lang, "experience"), styles)
    for item in content["experience"]["items"]:
        icon_markup = pdf_icon_markup(company_icon_path(item, build_dir / "company-icons"), 10)
        company_text = (
            pdf_link(item["company"], item["companyUrl"], style_colors["accent"])
            if item.get("companyUrl")
            else html.escape(item["company"])
        )
        story.append(pdf_markup(
            f"{icon_markup} {company_text} | <b>{pdf_text(experience_header_text(item))}</b>".strip(),
            styles["company"],
        ))
        story.append(pdf_paragraph(item["intro"], styles["body"]))
        for bullet in item["bullets"]:
            story.append(pdf_paragraph(f"• {bullet}", styles["bullet"]))
        story.append(pdf_markup(
            f"<b>{pdf_text(labels['stack'])}:</b> {pdf_text(', '.join(item['stack']))}",
            styles["meta"],
        ))
        story.append(Spacer(1, 1))

    pdf_section(story, download_section_title(lang, "education"), styles)
    for item in content["education"]["items"]:
        icon_markup = pdf_icon_markup(institution_icon_path(item, build_dir / "institution-icons"), 10)
        institution_text = (
            pdf_link(item["institution"], item["institutionUrl"], style_colors["accent"])
            if item.get("institutionUrl")
            else f"<b>{pdf_text(item['institution'])}</b>"
        )
        story.append(pdf_markup(
            f"{icon_markup} {institution_text} | {pdf_text(education_header_text(item))}".strip(),
            styles["body"],
        ))

    pdf_section(story, download_section_title(lang, "skills"), styles)
    skill_markup = "<br/>".join(
        f"<b>{pdf_text(title)}:</b> {pdf_text(values)}"
        for title, values in skill_group_segments(content["skills"]["groups"])
    )
    story.append(pdf_markup(skill_markup, styles["body"]))

    doc.build(story)


def copy_public_data(data_path: Path, output_root: Path) -> None:
    target_dir = output_root
    target_dir.mkdir(parents=True, exist_ok=True)
    data_target = target_dir / data_path.name
    if data_path.resolve() != data_target.resolve():
        shutil.copy2(data_path, data_target)
    schema_path = data_path.with_name("resume.schema.yaml")
    schema_target = target_dir / schema_path.name
    if schema_path.exists() and schema_path.resolve() != schema_target.resolve():
        shutil.copy2(schema_path, schema_target)


def contact_site_payload(contact: dict[str, str]) -> dict[str, str]:
    payload = {
        key: contact[key]
        for key in ("value", "icon", "iconDark")
        if key in contact
    }
    href = contact_href(contact)
    if href:
        payload["href"] = href
    return payload


def write_site_data(data: dict[str, Any], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    default_language = data["meta"]["defaultLanguage"]
    payload = {
        "contacts": {
            key: contact_site_payload(data["contacts"][key][default_language])
            for key in HERO_CONTACT_KEYS
            if key in data["contacts"]
        },
        "content": data["content"],
    }
    output_path.write_text(
        "window.JORQEN_RESUME_DATA = "
        + json.dumps(payload, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )


def generate_outputs(
    data_path: Path,
    data: dict[str, Any],
    output_root: Path,
    site_output: Path,
    build_dir: Path,
) -> None:
    copy_public_data(data_path, output_root)
    write_site_data(data, site_output)
    for lang in data["meta"]["languages"]:
        lang_dir = output_root / lang
        file_names = data["content"][lang]["resume"]["downloadNames"]
        generate_txt(data, lang, lang_dir / file_names["txt"])
        generate_docx(data, lang, lang_dir / file_names["docx"], build_dir)
        generate_pdf(data, lang, lang_dir / file_names["pdf"], build_dir)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data", type=Path, default=ROOT / "resume/resume.yaml")
    parser.add_argument("--output-root", type=Path, default=ROOT / "resume")
    parser.add_argument("--site-output", type=Path, default=ROOT / "assets/generated/resume-content.js")
    parser.add_argument("--build-dir", type=Path, default=ROOT / ".build/resume-assets")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = load_data(args.data)
    generate_outputs(args.data, data, args.output_root, args.site_output, args.build_dir)
    print("Resume outputs generated successfully.")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # pragma: no cover - keeps CI/local failures readable.
        print(f"Resume generation failed: {exc}", file=sys.stderr)
        raise
